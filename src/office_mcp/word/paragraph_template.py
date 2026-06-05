"""Generic Word template clone for paragraph-based templates.

Replaces paragraph text by index while preserving all formatting.
"""
import copy
from pathlib import Path
from docx import Document
from lxml import etree

_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_Wp = f'{{{_W}}}'


def clone_paragraph_template(template_path: str, output_path: str, replacements: dict) -> str:
    """Clone a paragraph-based template, replacing text by paragraph index.

    Args:
        template_path: path to the .docx template
        output_path: where to save the new .docx
        replacements: dict mapping paragraph_index -> new_text

    Returns:
        output_path as string
    """
    doc = Document(template_path)
    cell = None

    # If a table exists, operate on its first cell's paragraphs
    if doc.tables:
        cell = doc.tables[0].rows[0].cells[0] if doc.tables[0].rows else None

    # Get all <w:p> elements (from cell if present, else from body)
    if cell is not None:
        p_list = cell._tc.findall(f'{_Wp}p')
    else:
        body = doc.element.body
        p_list = body.findall(f'{_Wp}p')

    # Replace text in specified paragraphs
    for idx, new_text in replacements.items():
        idx = int(idx)
        if idx < len(p_list):
            p = p_list[idx]
            # Clear all <w:t> text nodes
            for r in p.findall(f'{_Wp}r'):
                t = r.find(f'{_Wp}t')
                if t is not None:
                    t.text = ''
            # Set text in first run's <w:t>, or create one
            rs = p.findall(f'{_Wp}r')
            if rs:
                t = rs[0].find(f'{_Wp}t')
                if t is not None:
                    t.text = new_text
                else:
                    etree.SubElement(rs[0], f'{_Wp}t').text = new_text
            else:
                r = etree.SubElement(p, f'{_Wp}r')
                etree.SubElement(r, f'{_Wp}t').text = new_text

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)
