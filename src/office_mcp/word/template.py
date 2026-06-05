"""Word document template cloning engine."""
# -*- coding: utf-8 -*-
import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt
from lxml import etree


def clone_word_template(
    template_path: str,
    output_path: str,
    replacements: dict,
    date_str: str = None,
) -> str:
    """
    Clone a Word template document and replace specified content.
    
    Args:
        template_path: Path to the template .docx file.
        output_path: Path to save the new .docx file.
        replacements: Dict mapping field names to new content.
                     Supported keys:
                     - "title": Document title (paragraph 1, bold centered)
                     - "main_topic": Main content/topic (table cell)
                     - "process_record": Full process record text (table cell)
                     - "date": Date string (overrides date_str)
        date_str: Date string to replace in the date cell (e.g. "2026年6月5日 星期五 下午").
                   Deprecated: use replacements["date"] instead.
    
    Returns:
        Path to the saved document.
    """
    doc = Document(template_path)
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # --- Replace title (paragraph 1, centered, bold) ---
    if "title" in replacements:
        _replace_title(doc, replacements["title"])

    # --- Replace date ---
    effective_date = replacements.get("date", date_str)
    if effective_date:
        _replace_date(doc, effective_date, ns_w)

    # --- Replace main topic in table ---
    if "main_topic" in replacements and doc.tables:
        _replace_main_topic(doc, replacements["main_topic"], ns_w)

    # --- Replace process record (full merged cell content) ---
    if "process_record" in replacements and doc.tables:
        _replace_process_record(doc, replacements["process_record"])

    # --- Save ---
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


def _replace_title(doc: Document, title: str):
    """Replace the centered bold title paragraph."""
    target_para = None
    for i, para in enumerate(doc.paragraphs):
        if i == 1 and para.alignment and para.alignment == 1:  # CENTER
            target_para = para
            break

    if target_para is None and len(doc.paragraphs) > 1:
        target_para = doc.paragraphs[1]

    if target_para is None:
        return

    # Clear ALL runs, then set the first run to the new title
    for run in target_para.runs:
        run.text = ''
    if target_para.runs:
        target_para.runs[0].text = title
        target_para.runs[0].bold = True
    else:
        run = target_para.add_run(title)
        run.bold = True


def _replace_date(doc: Document, date_str: str, ns_w: str):
    """Replace date in the first table's date cell."""
    if not doc.tables:
        return
    table = doc.tables[0]
    if len(table.rows) > 0 and len(table.rows[0].cells) > 1:
        cell = table.rows[0].cells[1]
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = date_str
        else:
            run = para.add_run(date_str)
            run.font.name = '宋体'
            run.font.size = Pt(12)


def _replace_main_topic(doc: Document, topic: str, ns_w: str):
    """Replace main topic in table row 2, col 1 (merged)."""
    if not doc.tables:
        return
    table = doc.tables[0]
    if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
        cell = table.rows[2].cells[1]
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = topic
        else:
            _clear_cell_text(cell, ns_w)
            para = cell.paragraphs[0]
            run = para.add_run(topic)
            run.font.name = '宋体'
            run.font.size = Pt(12)


def _replace_process_record(doc: Document, content: str):
    """
    Replace process record by directly manipulating XML text nodes.
    
    CORE PRINCIPLE: Preserve ALL <w:p> paragraph elements exactly as-is.
    Only replace text inside <w:t> nodes. This guarantees 100% format 
    consistency (spacing, indentation, line-spacing) with the template.
    
    Empty paragraphs in the template are kept as-is (they control spacing).
    If new content has more lines than template paragraphs, extra paragraphs
    are appended with the same pPr as the last template paragraph.
    """
    if not doc.tables:
        return
    table = doc.tables[0]
    if len(table.rows) <= 3:
        return

    cell = table.rows[3].cells[0]
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Get ALL <w:p> elements in the cell (preserve order)
    p_list = cell._tc.findall(f'{{{ns_w}}}p')
    
    if not p_list:
        return
    
    # Save the pPr of the LAST non-empty paragraph (for appending if needed)
    last_pPr = None
    for p_elem in reversed(p_list):
        pPr = p_elem.find(f'{{{ns_w}}}pPr')
        if pPr is not None:
            last_pPr = copy.deepcopy(pPr)
            break
    
    # Step 1: Clear ALL text from ALL <w:t> nodes in ALL paragraphs
    # (but KEEP the <w:p> and <w:pPr> elements!)
    for p_elem in p_list:
        for r_elem in p_elem.findall(f'{{{ns_w}}}r'):
            t_elem = r_elem.find(f'{{{ns_w}}}t')
            if t_elem is not None:
                t_elem.text = ''
    
    # Step 2: Fill new content into existing <w:r> elements line-by-line
    # Strategy: each line of new content goes into one <w:p> paragraph.
    # We reuse the FIRST <w:r> in the paragraph, and CLEAR other <w:r> elements.
    lines = content.strip().split('\n')
    
    for i, line in enumerate(lines):
        if i < len(p_list):
            p_elem = p_list[i]
            # Find the first <w:r> that has a <w:t>
            r_list = p_elem.findall(f'{{{ns_w}}}r')
            target_r = None
            for r_elem in r_list:
                t_elem = r_elem.find(f'{{{ns_w}}}t')
                if t_elem is not None:
                    target_r = r_elem
                    break
            
            if target_r is None:
                # No <w:t> found, create one
                t_elem = etree.SubElement(r_list[0] if r_list else p_elem, f'{{{ns_w}}}t')
                t_elem.text = line
            else:
                t_elem = target_r.find(f'{{{ns_w}}}t')
                t_elem.text = line
            
            # Remove extra <w:r> elements (keep only the first one with text)
            for r_elem in list(r_list):
                t_elem = r_elem.find(f'{{{ns_w}}}t')
                if t_elem is not None and t_elem.text == line:
                    continue  # Keep this one
                # Remove extra <w:r> (but not the one we just filled)
                if r_elem != target_r:
                    p_elem.remove(r_elem)
        else:
            # New content exceeds template length: append new <w:p>
            new_p = etree.SubElement(cell._tc, f'{{{ns_w}}}p')
            if last_pPr is not None:
                new_p.append(copy.deepcopy(last_pPr))
            new_r = etree.SubElement(new_p, f'{{{ns_w}}}r')
            new_t = etree.SubElement(new_r, f'{{{ns_w}}}t')
            new_t.text = line
    
    # Step 3: All remaining template paragraphs stay EMPTY (preserving spacing)
    # (nothing to do - we already cleared their text in Step 1)


def _clear_cell_text(cell, ns_w: str):
    """Clear all text content in a table cell via XML."""
    for p in cell._tc.findall(f'.//{{{ns_w}}}p'):
        for r in p.findall(f'.//{{{ns_w}}}r'):
            for t in r.findall(f'.//{{{ns_w}}}t'):
                t.text = ''


def analyze_template(template_path: str) -> dict:
    """
    Analyze a Word template and return its structure as a dict.
    Useful for understanding what fields can be replaced.
    
    Returns:
        {
            "paragraphs": [{"index": 0, "style": "...", "text": "..."}],
            "tables": [{"rows": N, "cols": M, "cells": [...]}],
            "replaceable_fields": ["title", "date", "main_topic", "process_record"]
        }
    """
    doc = Document(template_path)
    result = {
        "paragraphs": [],
        "tables": [],
        "replaceable_fields": []
    }

    for i, para in enumerate(doc.paragraphs):
        result["paragraphs"].append({
            "index": i,
            "style": para.style.name if para.style else "Normal",
            "alignment": str(para.alignment),
            "text": para.text[:100] if para.text else "",
        })

    for ti, table in enumerate(doc.tables):
        table_info = {
            "table_index": ti,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "sample_cells": []
        }
        for ri, row in enumerate(table.rows[:3]):  # first 3 rows
            for ci, cell in enumerate(row.cells[:3]):  # first 3 cols
                table_info["sample_cells"].append({
                    "row": ri, "col": ci,
                    "text": cell.text.strip()[:50]
                })
        result["tables"].append(table_info)

    # Infer replaceable fields based on content
    text = '\n'.join(p.text for p in doc.paragraphs)
    if doc.tables:
        for row in doc.tables[0].rows:
            for cell in row.cells:
                text += '\n' + cell.text
    
    if '教研活动记录' in text or '活动' in text:
        result["replaceable_fields"] = [
            "title", "date", "location", "host", 
            "recorder", "main_topic", "process_record", "summary"
        ]

    return result
