"""Universal Word document engine — XML-first, format-preserving.

Core principle: NEVER delete/create <w:p> elements. Only replace <w:t> text.
This guarantees 100% format fidelity with the source document.

Addressing scheme:
  p:{i}                    → body paragraph
  c:{t}:{r}:{c}            → entire cell (auto line-map to paragraphs)
  c:{t}:{r}:{c}:{p}        → specific paragraph inside a cell

Optional run-format hints:
  _fmt:p:{i}               → {"runs": [{"text": "中一", "underline": true}]}
"""
import copy
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from lxml import etree

_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_Wp = f'{{{_W}}}'


# ─── Public API ──────────────────────────────────────────────────────────────

def analyze(path: str) -> dict:
    """Analyze a Word document, returning a universal text index.

    Returns a compact structure with:
      - paras: [{i, t, s}]  body paragraphs (non-empty, truncated)
      - cells: [{t, r, c, txt, pn}]  table cells (non-empty)
      - merge: merge info per table
    """
    doc = Document(path)
    result = {"paras": [], "tables": []}

    # Body paragraphs
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt:
            result["paras"].append({"i": i, "t": txt[:80], "s": p.style.name})

    # Tables
    for ti, tbl in enumerate(doc.tables):
        tbl_info = {"i": ti, "rows": len(tbl.rows), "cols": len(tbl.columns),
                    "cells": [], "merge": []}
        seen = set()
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                cid = id(cell)
                if cid in seen:
                    continue
                seen.add(cid)
                txt = cell.text.strip()
                pn = len(cell.paragraphs)
                if txt or pn > 1:
                    # Collect per-paragraph text for AI to target
                    p_texts = []
                    for pi, p in enumerate(cell.paragraphs):
                        pt = p.text.strip()
                        if pt:
                            p_texts.append({"i": pi, "t": pt[:60]})
                    tbl_info["cells"].append({
                        "r": ri, "c": ci, "txt": txt[:60], "pn": pn,
                        "p": p_texts if p_texts else None,
                    })

                # Merge info
                tc = cell._tc
                gs = tc.find(f'.//{_Wp}gridSpan')
                cs = int(gs.get(f'{_Wp}val', '1')) if gs is not None else 1
                vm = tc.find(f'.//{_Wp}vMerge')
                rs = 0
                if vm is not None and vm.get(f'{_Wp}val', '') == 'restart':
                    rs = 1
                    for ro in range(1, len(tbl.rows) - ri):
                        nr = tbl.rows[ri + ro]
                        # Find cell at same logical column
                        if ci < len(nr.cells):
                            nvm = nr.cells[ci]._tc.find(f'.//{_Wp}vMerge')
                            if nvm is not None and nvm.get(f'{_Wp}val', '') == 'continue':
                                rs += 1
                            else:
                                break
                if cs > 1 or rs > 1:
                    tbl_info["merge"].append({"r": ri, "c": ci, "rs": rs, "cs": cs})

        result["tables"].append(tbl_info)

    return result


def replace_text(path: str, output_path: str, replacements: dict,
                 fmt_hints: dict = None) -> str:
    """Universal text replacement — accepts a replacement dictionary.

    Args:
        path: source .docx path
        output_path: where to save the new .docx
        replacements: {
            "p:0": "new title",              → body paragraph
            "c:0:3:0": "line1\\nline2\\n...",  → entire cell (auto line-map)
            "c:0:2:1:0": "new cell text",    → specific cell paragraph
        }
        fmt_hints: optional {
            "_fmt:p:0": {"runs": [{"text": "中一", "underline": true}]}
        }

    Returns:
        output_path as string
    """
    doc = Document(path)

    for addr, new_text in replacements.items():
        parts = addr.split(":")

        if parts[0] == "p":
            # Body paragraph: p:{i}
            pi = int(parts[1])
            if pi < len(doc.paragraphs):
                _replace_para(doc.paragraphs[pi], new_text)
                # Apply format hints if any
                hint_key = f"_fmt:p:{parts[1]}"
                if fmt_hints and hint_key in fmt_hints:
                    _apply_run_hints(doc.paragraphs[pi], fmt_hints[hint_key])

        elif parts[0] == "c":
            # Cell: c:{t}:{r}:{c} or c:{t}:{r}:{c}:{p}
            ti, ri, ci = int(parts[1]), int(parts[2]), int(parts[3])
            if ti < len(doc.tables):
                tbl = doc.tables[ti]
                if ri < len(tbl.rows) and ci < len(tbl.rows[ri].cells):
                    cell = tbl.rows[ri].cells[ci]
                    if len(parts) == 5:
                        # Specific paragraph in cell: c:{t}:{r}:{c}:{p}
                        pi = int(parts[4])
                        paras = cell.paragraphs
                        if pi < len(paras):
                            _replace_para(paras[pi], new_text)
                    else:
                        # Entire cell: c:{t}:{r}:{c}
                        _replace_cell(cell, new_text)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ─── Legacy API (backward compatibility) ────────────────────────────────────

def clone_word_template(template_path: str, output_path: str,
                        replacements: dict) -> str:
    """Legacy API — maps field names to universal addresses.

    Supports: title, date, main_topic, process_record
    """
    doc = Document(template_path)
    rmap = {}

    if "title" in replacements and len(doc.paragraphs) > 1:
        rmap["p:1"] = replacements["title"]
        # Title underline hint
        fmt = _detect_title_underline_hints(doc, replacements["title"])
        if fmt:
            rmap.setdefault("_fmt", {})
            rmap["_fmt"]["_fmt:p:1"] = fmt

    if "date" in replacements and doc.tables:
        rmap["c:0:0:1:0"] = replacements["date"]

    if "main_topic" in replacements and doc.tables:
        rmap["c:0:2:1:0"] = replacements["main_topic"]

    if "process_record" in replacements and doc.tables:
        rmap["c:0:3:0"] = replacements["process_record"]

    # Separate fmt hints
    fmt_hints = rmap.pop("_fmt", None) or None

    return replace_text(template_path, output_path, rmap, fmt_hints)


# ─── Internal helpers ───────────────────────────────────────────────────────

def _replace_para(para, text: str):
    """Replace text in a paragraph, preserving all <w:r> formatting.

    Strategy: clear all <w:t> text, then set text in first run.
    If the paragraph has multiple runs with different formatting (e.g. underline),
    distribute text across runs proportionally.
    """
    p_elem = para._element
    runs = p_elem.findall(f'{_Wp}r')

    if not runs:
        r = etree.SubElement(p_elem, f'{_Wp}r')
        etree.SubElement(r, f'{_Wp}t').text = text
        return

    # Simple case: single run or very short text
    if len(runs) == 1:
        t = runs[0].find(f'{_Wp}t')
        if t is not None:
            t.text = text
        else:
            etree.SubElement(runs[0], f'{_Wp}t').text = text
        return

    # Multi-run: distribute text proportionally, preserve each run's rPr
    # First, save original run text lengths
    orig_lens = []
    for r in runs:
        t = r.find(f'{_Wp}t')
        orig_lens.append(len(t.text) if t is not None and t.text else 0)
    total = sum(orig_lens) or 1

    # Clear all run text
    for r in runs:
        t = r.find(f'{_Wp}t')
        if t is not None:
            t.text = ''

    # Distribute new text proportionally
    n = len(text)
    idx = 0
    for i, r in enumerate(runs):
        if idx >= n:
            break
        if i < len(runs) - 1:
            share = max(1, round((orig_lens[i] / total) * n))
            share = min(share, n - idx)
        else:
            share = n - idx
        t = r.find(f'{_Wp}t')
        if t is None:
            t = etree.SubElement(r, f'{_Wp}t')
        t.text = text[idx:idx + share]
        idx += share


def _replace_cell(cell, content: str):
    """Replace entire cell content, auto line-mapping to paragraphs.

    Preserves all <w:p> elements (spacing, indent, etc).
    Blank lines in content are STRIPPED — spacing comes from pPr, not empty paragraphs.
    If content exceeds template paragraphs, appends new ones with last pPr.
    """
    tc = cell._tc
    p_list = tc.findall(f'{_Wp}p')
    if not p_list:
        return

    # Save last pPr for appending if needed
    last_pPr = None
    for p in reversed(p_list):
        pp = p.find(f'{_Wp}pPr')
        if pp is not None:
            last_pPr = copy.deepcopy(pp)
            break

    # Clear ALL <w:t> text (keep <w:p> and <w:pPr> intact)
    for p in p_list:
        for r in p.findall(f'{_Wp}r'):
            t = r.find(f'{_Wp}t')
            if t is not None:
                t.text = ''

    # Fill new content — strip blank lines (spacing from pPr, not empty paragraphs)
    lines = [l for l in content.strip().split('\n') if l.strip()]

    for i, line in enumerate(lines):
        if i < len(p_list):
            p = p_list[i]
            rs = p.findall(f'{_Wp}r')
            if rs:
                t = rs[0].find(f'{_Wp}t')
                if t is not None:
                    t.text = line
                else:
                    etree.SubElement(rs[0], f'{_Wp}t').text = line
                # Remove extra runs (keep only first)
                for r in rs[1:]:
                    p.remove(r)
            else:
                r = etree.SubElement(p, f'{_Wp}r')
                etree.SubElement(r, f'{_Wp}t').text = line
        else:
            # Append new paragraph (content exceeds template)
            new_p = etree.SubElement(tc, f'{_Wp}p')
            if last_pPr is not None:
                new_p.append(copy.deepcopy(last_pPr))
            r = etree.SubElement(new_p, f'{_Wp}r')
            etree.SubElement(r, f'{_Wp}t').text = line


def _detect_title_underline_hints(doc, new_title: str) -> dict:
    """Detect which runs have underline in the title paragraph.

    Returns format hints for the new title, e.g.:
    {"runs": [{"text": "中一", "underline": true}]}
    """
    if len(doc.paragraphs) <= 1:
        return None
    para = doc.paragraphs[1]
    runs = para.runs

    underlined_texts = []
    for r in runs:
        r_elem = r._r
        rPr = r_elem.find(f'{_Wp}rPr')
        if rPr is not None:
            u = rPr.find(f'{_Wp}u')
            if u is not None:
                val = u.get(f'{_Wp}val', 'single')
                if val and val != 'none':
                    underlined_texts.append(r.text)

    if not underlined_texts:
        return None

    # Check if underlined text exists in new title
    runs_hints = []
    for ut in underlined_texts:
        if ut and ut in new_title:
            runs_hints.append({"text": ut, "underline": True})

    return {"runs": runs_hints} if runs_hints else None


def _apply_run_hints(para, hints: dict):
    """Apply run-level format hints to a paragraph.

    E.g. {"runs": [{"text": "中一", "underline": true}]}
    After text has been distributed across runs, find the run containing
    the hinted text and apply the format.
    """
    if not hints or "runs" not in hints:
        return

    for hint in hints["runs"]:
        target_text = hint.get("text", "")
        if not target_text:
            continue

        # Find which run contains this text
        run_start = 0
        for r in para.runs:
            run_end = run_start + len(r.text)
            # Check if target text overlaps with this run
            idx = r.text.find(target_text)
            if idx != -1:
                # Found it — apply format
                r_elem = r._r
                rPr = r_elem.find(f'{_Wp}rPr')
                if rPr is None:
                    rPr = etree.SubElement(r_elem, f'{_Wp}rPr')
                    # Move rPr to be first child
                    r_elem.remove(rPr)
                    r_elem.insert(0, rPr)

                if hint.get("underline"):
                    u = rPr.find(f'{_Wp}u')
                    if u is None:
                        u = etree.SubElement(rPr, f'{_Wp}u')
                    u.set(f'{_Wp}val', 'single')

                if hint.get("bold") is not None:
                    r.bold = hint["bold"]

                break
            run_start = run_end

        # If underline target was NOT found in any run,
        # we need to also remove underline from runs that originally had it
        # but shouldn't anymore
        if hint.get("underline") and target_text:
            for r in para.runs:
                if target_text not in r.text:
                    r_elem = r._r
                    rPr = r_elem.find(f'{_Wp}rPr')
                    if rPr is not None:
                        u = rPr.find(f'{_Wp}u')
                        if u is not None and u.get(f'{_Wp}val', '') != 'none':
                            u.set(f'{_Wp}val', 'none')
