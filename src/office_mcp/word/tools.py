"""Word document MCP tools."""

from fastmcp import FastMCP, Context
from pathlib import Path
from .template import clone_word_template, analyze_template
from .table_parser import parse_table_format, analyze_table_merge, format_table_analysis

# Initialize FastMCP server
mcp = FastMCP("office-word")


@mcp.tool()
def word_clone_template(
    template_path: str,
    output_path: str,
    title: str = None,
    date: str = None,
    main_topic: str = None,
    process_record: str = None,
    process_record_file: str = None,
    ctx: Context = None,
) -> dict:
    """
    Clone a Word template document and replace specified fields.
    
    Perfect for reusing document templates (e.g. meeting minutes, 
    lesson plans) with new content while preserving all formatting.
    
    Args:
        template_path: Path to the .docx template file.
        output_path: Where to save the new .docx file.
        title: (Optional) New title text (replaces centered bold title).
        date: (Optional) New date string (replaces date in table).
        main_topic: (Optional) New main topic (replaces topic in table).
        process_record: (Optional) Full process record text (replaces merged cell content).
        process_record_file: (Optional) Path to file containing process record text.
            If both process_record and process_record_file are provided,
            process_record_file takes priority.
    
    Returns:
        Dict with success status and output path.
    """
    try:
        replacements = {}
        if title:
            replacements["title"] = title
        if date:
            replacements["date"] = date
        if main_topic:
            replacements["main_topic"] = main_topic
        
        # Handle process_record from file or direct input
        final_process_record = process_record
        if process_record_file:
            try:
                with open(process_record_file, 'r', encoding='utf-8') as f:
                    final_process_record = f.read()
            except Exception as file_e:
                return {"success": False, "error": f"Failed to read process_record_file: {str(file_e)}"}
        
        if final_process_record:
            replacements["process_record"] = final_process_record

        result = clone_word_template(
            template_path=template_path,
            output_path=output_path,
            replacements=replacements,
        )
        return {"success": True, "output_path": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def word_analyze_template(
    template_path: str,
    ctx: Context = None,
) -> dict:
    """
    Analyze a Word template and return its structure.
    
    Use this to understand what fields can be replaced in a template
    before calling word_clone_template.
    
    Args:
        template_path: Path to the .docx template file.
    
    Returns:
        Dict describing the document structure and replaceable fields.
        Also includes table format information in 'table_format' field.
    """
    try:
        from docx import Document
        
        result = analyze_template(template_path)
        
        # NEW: Add table format analysis
        table_format = []
        try:
            doc = Document(template_path)
            for table in doc.tables:
                format_info = format_table_analysis(table)
                table_format.append(format_info)
        except Exception as tf_e:
            # Don't fail the whole request if table format analysis fails
            table_format = {"error": str(tf_e)}
        
        return {
            "success": True,
            "structure": result,
            "table_format": table_format,  # NEW field
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def word_create_document(
    output_path: str,
    title: str = None,
    content: str = None,
    ctx: Context = None,
) -> dict:
    """
    Create a new Word document from scratch.
    
    Args:
        output_path: Where to save the .docx file.
        title: (Optional) Document title.
        content: (Optional) Document content (use \\n to separate paragraphs).
    
    Returns:
        Dict with success status and output path.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if content:
            for line in content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)
        
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return {"success": True, "output_path": output_path}
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def word_read_document(
    document_path: str,
    ctx: Context = None,
) -> dict:
    """
    Read and extract text content from a Word document.
    
    Args:
        document_path: Path to the .docx file.
    
    Returns:
        Dict with document text content and structure info.
        Also includes table cell format information in 'tables[].cell_format'.
    """
    try:
        from docx import Document
        
        doc = Document(document_path)
        
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append({
                    "index": i,
                    "style": para.style.name if para.style else "Normal",
                    "text": para.text,
                })
        
        tables = []
        for ti, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            # NEW: Add cell format information
            try:
                cell_format = format_table_analysis(table)
            except Exception as tf_e:
                cell_format = {"error": str(tf_e)}
            
            tables.append({
                "table_index": ti,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data,  # Original field, untouched
                "cell_format": cell_format,  # NEW field
            })
        
        return {
            "success": True,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "paragraphs": paragraphs,
            "tables": tables,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}
