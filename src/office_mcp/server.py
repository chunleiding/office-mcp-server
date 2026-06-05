#!/usr/bin/env python3
"""Office MCP Server - Main entry point.

Supports Word/Excel/PPT document automation via MCP tools.
"""

import argparse
import sys
from pathlib import Path

from fastmcp import FastMCP

# =============================================================================
# Build the combined server
# =============================================================================

mcp = FastMCP("office-mcp-server")


# =============================================================================
# Utility: normalize special chars that can break JSON-RPC encoding
# =============================================================================

def _normalize_text(text: str) -> str:
    """Normalize special characters that break JSON-RPC encoding.

    Chinese quotation marks are the most common culprit — when the
    MCP client sends them raw (without json.dumps-level escaping),
    they can corrupt the JSON-RPC message and cause "root: must be
    object" / "Expecting value" style errors.

    Normalizing them to ASCII equivalents is lossless for document
    content (the meaning is preserved) and makes JSON-RPC rock-solid.

    Args:
        text: Input text that may contain problematic characters.

    Returns:
        Normalized text with special characters replaced.
    """
    text = text.replace("\u201c", '"')  # " (LEFT DOUBLE QUOTATION MARK) -> "
    text = text.replace("\u201d", '"')  # " (RIGHT DOUBLE QUOTATION MARK) -> "
    text = text.replace("\u2018", "'")  # ' (LEFT SINGLE QUOTATION MARK) -> '
    text = text.replace("\u2019", "'")  # ' (RIGHT SINGLE QUOTATION MARK) -> '
    return text


# =============================================================================
# Word Tools
# =============================================================================

@mcp.tool()
def word_clone_template(
    template_path: str,
    output_path: str,
    title: str = None,
    date: str = None,
    main_topic: str = None,
    process_record: str = None,
    process_record_file: str = None,
) -> dict:
    """Clone a Word template document and replace specified fields.

    Perfect for reusing document templates (e.g. meeting minutes,
    lesson plans) with new content while preserving all formatting.

    Args:
        template_path: Path to the .docx template file.
        output_path: Where to save the new .docx file.
        title: (Optional) New title text (replaces centered bold title).
        date: (Optional) New date string (replaces date in table).
        main_topic: (Optional) New main topic (replaces topic in table).
        process_record: (Optional) Full process record text.
            For short content (< 500 chars, no special chars).
        process_record_file: (Optional) Path to a .txt file containing
            the process record content. **Recommended for long text or
            text with special characters.**  If both are provided,
            process_record_file takes priority.

    Returns:
        Dict with success status and output path.
    """
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from office_mcp.word.template import clone_word_template as do_clone

    try:
        replacements = {}
        if title:
            replacements["title"] = title
        if date:
            replacements["date"] = date
        if main_topic:
            replacements["main_topic"] = main_topic

        # Resolve process_record: file takes priority
        final = process_record
        if process_record_file:
            try:
                with open(process_record_file, "r", encoding="utf-8") as f:
                    final = f.read()
            except OSError as exc:
                return {
                    "success": False,
                    "error": f"Cannot read process_record_file: {exc}",
                }

        if final:
            # Normalize special chars so the content survives JSON-RPC
            final = _normalize_text(final)
            replacements["process_record"] = final

        result = do_clone(
            template_path=template_path,
            output_path=output_path,
            replacements=replacements,
        )
        return {"success": True, "output_path": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def word_analyze_template(
    template_path: str,
) -> dict:
    """Analyze a Word template and return its structure.

    Use this to understand what fields can be replaced in a template
    before calling word_clone_template.

    Args:
        template_path: Path to the .docx template file.

    Returns:
        Dict describing the document structure, replaceable fields,
        and table format information.
    """
    sys.path.insert(0, _SRC_DIR)
    from office_mcp.word.template import analyze_template

    try:
        result = analyze_template(template_path)

        # Attach table format analysis
        table_format = []
        try:
            from docx import Document
            from office_mcp.word.table_parser import format_table_analysis

            doc = Document(template_path)
            for table in doc.tables:
                table_format.append(format_table_analysis(table))
        except Exception:
            table_format = []

        return {
            "success": True,
            "structure": result,
            "table_format": table_format,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def word_create_document(
    output_path: str,
    title: str = None,
    content: str = None,
) -> dict:
    """Create a new Word document from scratch.

    Args:
        output_path: Where to save the .docx file.
        title: (Optional) Document title.
        content: (Optional) Document content (use \\n to separate paragraphs).

    Returns:
        Dict with success status and output path.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()

        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if content:
            for line in content.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(12)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return {"success": True, "output_path": output_path}
    except Exception as e:
        return {"success": False, "error": str(e)}


@mcp.tool()
def word_read_document(
    document_path: str,
) -> dict:
    """Read and extract text content from a Word document.

    Args:
        document_path: Path to the .docx file.

    Returns:
        Dict with paragraphs, tables, and cell-level format info.
    """
    sys.path.insert(0, _SRC_DIR)
    try:
        from docx import Document
        from office_mcp.word.table_parser import format_table_analysis

        doc = Document(document_path)

        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append({
                    "index": i,
                    "style": para.style.name if para.style else "Normal",
                    "text": para.text,
                })

        tables = []
        for ti, table in enumerate(doc.tables):
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]

            try:
                cell_format = format_table_analysis(table)
            except Exception:
                cell_format = {}

            tables.append({
                "table_index": ti,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data,          # original text-only field
                "cell_format": cell_format,  # NEW: full formatting info
            })

        return {
            "success": True,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "paragraphs": paragraphs,
            "tables": tables,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# =============================================================================
# Main
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Office MCP Server")
    parser.add_argument(
        "--transport",
        choices=["stdio", "sse"],
        default="stdio",
        help="Transport protocol (default: stdio)",
    )
    parser.add_argument(
        "--port",
        type=int,
        default=8000,
        help="Port for SSE transport (default: 8000)",
    )
    args = parser.parse_args()

    print(
        f"🚀 Starting Office MCP Server (transport={args.transport})",
        file=sys.stderr,
    )

    if args.transport == "sse":
        mcp.run(transport="sse", host="127.0.0.1", port=args.port)
    else:
        mcp.run(transport="stdio")


if __name__ == "__main__":
    main()
